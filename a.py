# auto-py-to-exe para compilar este script
# Asegúrate de tener instaladas las librerías necesarias
# pyinstaller --onefile --add-data "archivo.txt:." tu_script.py para linux


import os
import openpyxl
from tkinter import Tk, Label, Button, Text, Scrollbar, END, filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from google import genai

from docx.shared import Cm
import pandas as pd



def leer_datos_desde_excel(ruta):
    wb = openpyxl.load_workbook(ruta, data_only=True)
    global hoja
    hoja = wb.active
    datos = []

    fila = 23
    while True:
        # Lee solo las columnas C (3) hasta K (11)
        valores = [hoja.cell(row=fila, column=col).value for col in range(3, 12)]
        # Si la celda de la columna C está vacía, termina el bucle
        if hoja.cell(row=fila, column=3).value is None:
            break
        datos.append(valores)
        fila += 1

    return datos


def extraer_ponentes_rango_excel(ruta_archivo):
    wb = openpyxl.load_workbook(ruta_archivo, data_only=True)
    hoja = wb['FICHA ECONÓMICA']
    ponentes = {}

    # Leer encabezados desde la fila 22 (índice 21)
    encabezados = [hoja.cell(row=22, column=col).value for col in range(1, 12)]
    # Asumimos que los datos empiezan en la fila 23
    fila = 23
    while True:
        dni = hoja.cell(row=fila, column=4).value
        nombre = hoja.cell(row=fila, column=3).value
        concepto = hoja.cell(row=fila, column=5).value
        tipo_intervencion = hoja.cell(row=fila, column=6).value
        unidades = hoja.cell(row=fila, column=7).value
        importe = hoja.cell(row=fila, column=11).value

        # Si no hay DNI, terminamos
        if dni is None:
            break

        nombre = nombre.strip() if isinstance(nombre, str) else nombre
        partida = {
            'Concepto': concepto,
            'Tipo_Intervencion': tipo_intervencion,
            'Unidades': unidades,
            'Importe': importe
        }
        if dni not in ponentes:
            ponentes[dni] = {'Nombre': nombre, 'Partidas': []}
        ponentes[dni]['Partidas'].append(partida)
        fila += 1

    return str(ponentes)

    

def generar_documento(datos):
    doc = Document()
    # Establecer el borde superior (margen superior) a 0.6 inch
    section = doc.sections[0]
    section.top_margin = Cm(1)

    imagen_path = "a.png"
    doc.add_picture(imagen_path, width=Cm(15.0))  # Ajusta el ancho de la imagen según sea necesario
    # doc.add_paragraph("\n")

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(14)

    # Encabezado
    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.LEFT
    encabezado.add_run("\nJORDI MARTÍ GUIU, SUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT \n \n")
    

    # Crear una tabla con dos columnas: izquierda (estrecha), derecha (ancha)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False



    # Añadir contenido a la primera fila
    fila = tabla.rows[0]
    fila.cells[0].width = Cm(3)  # Ancho de la columna izquierda
    fila.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    run = fila.cells[0].paragraphs[0].add_run("DESIGNA")
    run.bold = True
    
    dni = str(datos[0][1])
    if datos[0][8] == 300:
        pagament_text = " setmanes a 300 euros/setmana"
    else:
        pagament_text = f" hores a {str(datos[0][7]).lower()} euros/hora"
    cadena1 = str(datos[0][5]).lower() + pagament_text +" en concepte de "+ str(datos[0][4]).lower()
    if dni == datos[1][1]: 
        if datos[1][8] == 300:
            pagament_text = " setmanes a 300 euros/setmana"
        else:
            pagament_text = f" hores a {str(datos[1][7]).lower()} euros/hora"
        cadena2 = ", "+str(datos[1][5]).lower() + pagament_text +" en concepte de "+ str(datos[1][4]).lower()
        if dni == datos[2][1]:
            if datos[2][8] == 300:
                pagament_text = " setmanes a 300 euros/setmana"
            else:
                pagament_text = f" hores a {str(datos[2][7]).lower()} euros/hora"
            cadena3 = ", "+str(datos[2][5]).lower() + pagament_text +" en concepte de "+ str(datos[2][4]).lower()
    

    cuerpo = (
        str(datos[0][0]).capitalize() + " amb nif: " + dni.lower() + "\n\n" + 
        str(datos[0][2]).lower() + ", perquè impartisca " + cadena1 + cadena2 + cadena3 + " del curs amb les següents dades:\n\n"
        "Codi: " + str(hoja.cell(row=13, column=4).value if hoja.cell(row=13, column=4).data_type != 'f' else hoja.cell(row=13, column=4).value).upper() + "\n\n"
        "Títol: " + str(hoja.cell(row=14, column=4).value if hoja.cell(row=14, column=4).data_type != 'f' else hoja.cell(row=14, column=4).value) + "\n\n"
        "data de realització: de " + str(hoja.cell(row=16, column=4).value if hoja.cell(row=16, column=4).data_type != 'f' else hoja.cell(row=16, column=4).value).lower() + "\n\n"
        "lloc de realització: " + str(hoja.cell(row=17, column=4).value if hoja.cell(row=17, column=4).data_type != 'f' else hoja.cell(row=17, column=4).value) + "\n\n"
        "caldrà pagar-li en concepte d´assistències la quantitat de " + str(datos[0][8]) + " euros (" + numero_a_letras(datos[0][8]).lower() + "), per l’aplicació 233.02"
    )

    fila.cells[1].width = Cm(13)  # Ancho de la columna izquierda


    fila.cells[1].text = cuerpo
    # doc.add_paragraph(cuerpo)


    doc.save("resultado.docx")
    messagebox.showinfo("Documento generado", "✅ Se ha creado 'resultado.docx' correctamente.")

def generar_documento_skills(datos):
    doc = Document()
    # Establecer el borde superior (margen superior) a 0.6 inch
    section = doc.sections[0]
    section.top_margin = Cm(1)



    #     for fila in datos:        doc.add_paragraph(', '.join([str(v) if v is not None else '' for v in fila]))

    imagen_path = "a.png"
    doc.add_picture(imagen_path, width=Cm(15.0))  # Ajusta el ancho de la imagen según sea necesario
    # doc.add_paragraph("\n")

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(14)

    # Encabezado
    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.LEFT
    encabezado.add_run("\n JORDI MARTÍ GUIU, SUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT \n \n")
    

    # Crear una tabla con dos columnas: izquierda (estrecha), derecha (ancha)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False



    # Añadir contenido a la primera fila
    fila = tabla.rows[0]
    fila.cells[0].width = Cm(3)  # Ancho de la columna izquierda
    fila.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    run = fila.cells[0].paragraphs[0].add_run("DESIGNA")
    run.bold = True
    

    # Puedes añadir más filas si lo necesitas, por ejemplo:
    # for fila_datos in datos:
    #     nueva_fila = tabla.add_row().cells
    #     nueva_fila[0].text = str(fila_datos[0])
    #     nueva_fila[1].text = str(fila_datos[1])


    # doc.add_paragraph("\n\nDESIGNA\n", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    dni = str(datos[0][1])
 
    # messagebox.showinfo("DNI", f"El DNI es: {type(dni)}")
    # Cuerpo principal
    if datos[0][8] == 300:
        pagament_text = " setmanes a 300 euros/setmana"
    else:
        pagament_text = f" hores a {str(datos[0][8]).lower()} euros/hora"

    cuerpo = (
        str(datos[0][0]).capitalize() + " amb nif: " + dni.lower() + "\n\n"
        "funcionari, perquè impartisca " + str(datos[0][6]).lower() + pagament_text + " del curs amb les següents dades:\n\n"
        "Codi: " + str(hoja.cell(row=13, column=4).value if hoja.cell(row=13, column=4).data_type != 'f' else hoja.cell(row=13, column=4).value).upper() + "\n\n"
        "Títol: " + str(hoja.cell(row=14, column=4).value if hoja.cell(row=14, column=4).data_type != 'f' else hoja.cell(row=14, column=4).value) + "\n\n"
        "data de realització: de " + str(hoja.cell(row=16, column=4).value if hoja.cell(row=16, column=4).data_type != 'f' else hoja.cell(row=16, column=4).value).lower() + "\n\n"
        "lloc de realització: " + str(hoja.cell(row=17, column=4).value if hoja.cell(row=17, column=4).data_type != 'f' else hoja.cell(row=17, column=4).value) + "\n\n"
        "caldrà pagar-li en concepte d´assistències la quantitat de " + str(datos[0][8]) + " euros (" + numero_a_letras(datos[0][8]).lower() + "), per l’aplicació 233.02"
    )

    fila.cells[1].width = Cm(13)  # Ancho de la columna izquierda


    fila.cells[1].text = cuerpo
    # doc.add_paragraph(cuerpo)


    doc.save("resultado.docx")
    messagebox.showinfo("Documento generado", "✅ Se ha creado 'resultado.docx' correctamente.")



def mostrar_datos_en_texto(datos, texto_widget):
    texto_widget.delete(1.0, END)
    for fila in datos:
        texto_widget.insert(END, f"{fila}\n")

def buscar_archivo(texto_widget):
    archivo = filedialog.askopenfilename(
        title="Selecciona archivo Excel",
        filetypes=[("Excel files", "*.xlsx")]
    )
    if archivo:
        try:
            #datos = leer_datos_desde_excel(archivo)
            datos = extraer_ponentes_rango_excel(archivo)
            mostrar_datos_en_texto(datos, texto_widget)
            Button(root, text="Generar DESIGNA FSE", command=lambda: generar_documento(datos)).pack(pady=10)
            Button(root, text="Generar DESIGNA SKILLS", command=lambda: generar_documento_skills(datos)).pack(pady=10)
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error: {e}")

def numero_a_letras(numero):
    client = genai.Client(api_key=cargar_api_key())
    response = client.models.generate_content(
        model="gemini-2.5-flash", contents=f"Convertix en valencià el número {numero} a lletres. Dona'm només la resposta sense cap explicació addicional."
    )
    return response.text

# 🖥️ Interfaz principal

import tkinter.simpledialog

def guardar_api_key(api_key):
    with open("api.txt", "w") as f:
        f.write(api_key)

def cargar_api_key():
    if os.path.exists("api.txt"):
        with open("api.txt", "r") as f:
            return f.read().strip()
    return ""

def establecer_api_key():
    api_key_actual = cargar_api_key()
    api_key = tkinter.simpledialog.askstring("API Key", "Introduce la API Key:", initialvalue=api_key_actual)
    if api_key:
        guardar_api_key(api_key)
        messagebox.showinfo("API Key", "API Key guardada correctamente.")

root = Tk()
root.title("Lector de Excel y Generador de Word")
root.geometry("800x600")

# Menú
menu_bar = tkinter.Menu(root)
config_menu = tkinter.Menu(menu_bar, tearoff=0)
config_menu.add_command(label="Establecer API Key", command=establecer_api_key)
menu_bar.add_cascade(label="Configuración", menu=config_menu)
root.config(menu=menu_bar)

Label(root, text="Pulsa el botón para seleccionar un archivo Excel (.xlsx)").pack(pady=10)
Button(root, text="Buscar archivo Excel", command=lambda: buscar_archivo(texto)).pack(pady=5)

texto = Text(root, wrap='none', width=100, height=25)
texto.pack()

scroll = Scrollbar(root, command=texto.yview)
scroll.pack(side='right', fill='y')
texto.config(yscrollcommand=scroll.set)

root.mainloop()