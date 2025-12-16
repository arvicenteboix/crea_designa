from flask import Flask, render_template, request, send_file, flash, redirect, url_for
import io
import os
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL  # ✅ CORREGIDO
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import pandas as pd
from openpyxl import load_workbook
import zipfile
import re

app = Flask(__name__)
app.secret_key = 'cefIRE-designas-2025'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'excel_file' not in request.files:
            flash('No se ha seleccionado archivo Excel', 'error')
            return render_template('index.html')
        
        file = request.files['excel_file']
        if file.filename == '':
            flash('No se ha seleccionado archivo Excel', 'error')
            return render_template('index.html')
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            # Procesar Excel (adaptado para funcionar sin tkinter)
            jsondata, identificativos = procesar_excel_flask(filepath)
            if identificativos == -1:
                flash('Error en los datos identificativos. Revisa las fechas.', 'error')
                return render_template('index.html')
            
            documentos_generados = []
            for persona in jsondata:
                try:
                    doc_data = generar_documento_persona(persona, identificativos)
                    documentos_generados.append({
                        'nombre': persona.get('Nombre', 'Sin nombre'),
                        'doc': doc_data
                    })
                except Exception as e:
                    flash(f'Error generando documento para {persona.get("Nombre", "Desconocido")}: {str(e)}', 'error')
            
            if documentos_generados:
                return generar_zip_response(documentos_generados, filename)
            else:
                flash('No se pudieron generar documentos', 'error')
                
        except Exception as e:
            flash(f'Error procesando Excel: {str(e)}', 'error')
    
    return render_template('index.html')

def procesar_excel_flask(filepath):
    """Procesar Excel sin dependencias de tkinter"""
    # Cargar workbook
    wb = load_workbook(filepath)
    
    # Buscar hojas relevantes
    hoja_movimientos = None
    hoja_personas = None
    hoja_identificativos = None
    
    for sheet in wb.sheetnames:
        if 'movimiento' in sheet.lower() or 'movs' in sheet.lower():
            hoja_movimientos = wb[sheet]
        elif 'persona' in sheet.lower():
            hoja_personas = wb[sheet]
        elif 'identificati' in sheet.lower() or 'curso' in sheet.lower():
            hoja_identificativos = wb[sheet]
    
    if not hoja_movimientos:
        raise ValueError("No se encontró hoja de movimientos")
    
    # Procesar movimientos (simplificado)
    movimientos_data = []
    for row in hoja_movimientos.iter_rows(min_row=2, values_only=True):
        if row[0]:  # Si hay datos
            mov = {
                'NOMBRE': row[0] if row[0] else '',
                'DNI': row[1] if row[1] else '',
                'TIPO DE INTERVENCIÓN': row[2] if row[2] else '',
                'UNIDADES/UNITATS': row[3] if row[3] else 0,
                'TARIFICACIÓN APLICADA': row[4] if row[4] else '',
                'IMPORTE / IMPORT': row[5] if row[5] else 0,
                'JURÍDICO': row[6] if row[6] else ''
            }
            movimientos_data.append(mov)
    
    # Agrupar por persona
    personas = {}
    for mov in movimientos_data:
        dni = str(mov['DNI']).strip()
        nombre = mov['NOMBRE'].strip()
        if dni not in personas:
            personas[dni] = {
                'Nombre': nombre,
                'DNI': dni,
                'Movimientos': []
            }
        personas[dni]['Movimientos'].append(mov)
    
    jsondata = list(personas.values())
    
    # Identificativos (simplificado)
    identificativos = {
        'CDIGO_EDICIN': '2025-001',
        'TTULO_ACCIN_FORMATIVA': 'CURSO EJEMPLO',
        'FECHAS_REALIZACION': '13/12/2025',
        'MODALIDAD/MODALITAT': 'PRESENCIAL'
    }
    
    return jsondata, identificativos

def generar_documento_persona(datos, identificativos):
    """Genera documento Word corregido"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Encabezado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MARTÍ GUIU")
    run.bold = True
    run.font.size = Pt(14)
    p.add_run("\nSUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT")
    
    # Tabla principal
    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False
    fila = tabla.rows[0]
    
    # Columna izquierda: DESIGNA
    cell_left = fila.cells[0]
    cell_left.width = Cm(3)
    p_left = cell_left.paragraphs[0]
    run_designa = p_left.add_run("DESIGNA")
    run_designa.bold = True
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    # Columna derecha: Contenido
    cell_right = fila.cells[1]
    cell_right.width = Cm(13)
    p_right = cell_right.paragraphs[0]
    
    # Cuerpo del documento
    add_normal_text(p_right, f"{datos['Nombre']} amb NIF {datos['DNI']}")
    add_normal_text(p_right, "Designat/da per aquesta Subdirecció, ha impartit satisfactòriament")
    
    # Datos del curso
    if identificativos:
        add_bold_text(p_right, "Codi: ")
        add_normal_text(p_right, identificativos.get('CDIGO_EDICIN', ''))
        add_bold_text(p_right, "Títol: ")
        add_normal_text(p_right, identificativos.get('TTULO_ACCIN_FORMATIVA', ''))
    
    # Total importe
    total_importe = sum(float(mov.get('IMPORTE / IMPORT', 0) or 0) 
                       for mov in datos.get('Movimientos', []))
    add_normal_text(p_right, f"cal fer-li el pagament corresponent per un total de {total_importe} euros.")
    
    docname = f"DESIGNA_{datos['Nombre'].replace(' ', '_')}.docx"
    
    # Guardar en buffer
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return {'name': docname, 'file': bio}

def add_bold_text(paragraph, text):
    """Añade texto en negrita"""
    run = paragraph.add_run(text)
    run.bold = True

def add_normal_text(paragraph, text):
    """Añade texto normal"""
    run = paragraph.add_run(text + " ")
    run.font.size = Pt(11)

def generar_zip_response(documentos, excel_name):
    """Genera ZIP con documentos"""
    bio = io.BytesIO()
    with zipfile.ZipFile(bio, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for doc in documentos:
            zipf.writestr(doc['name'], doc['file'].getvalue())
    bio.seek(0)
    return send_file(
        bio,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'designas_{excel_name.rsplit(".",1)[0]}.zip'
    )

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
