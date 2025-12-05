# auto-py-to-exe para compilar este script
# Asegúrate de tener instaladas las librerías necesarias
# pyinstaller --onefile --add-data "archivo.txt:." tu_script.py para linux

import os
import json
import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from tkinter import messagebox
from num2words import num2words
import re
from datetime import datetime
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import sys
from tkinter import ttk

from tkcalendar import DateEntry

version = "v1.0.9"


# from docx2pdf import convert

def resource_path(relative_path):
        if hasattr(sys, '_MEIPASS'):
            return os.path.join(sys._MEIPASS, relative_path)
        return os.path.join(os.path.abspath('.'), relative_path)

def find_excel_file(status_label):
    status_label.config(text="Buscando archivo Excel...")

    
    if sys.platform == "darwin":  # macOS

        file_path = filedialog.askopenfilename(
            title="Selecciona el archivo Excel",
            filetypes=[("Archivos Excel", "*.xlsx"), ("Todos los archivos", "*.*")]
        )
        if file_path:
            status_label.config(text=f"Archivo seleccionado: {os.path.basename(file_path)}")
            return file_path
        else:
            status_label.config(text="Búsqueda cancelada por el usuario.")
            return None
    else:
        files = [f for f in os.listdir('.') if f.lower().endswith(('.xlsx'))]
        if files:
            status_label.config(text=f"Archivo encontrado: {files[0]}")
            return files[0]
        else:
            status_label.config(text="No se encontró ningún archivo Excel.")
            return None


def normaliza_fechas_realizacion(fecha_str):
    """
    Normaliza diferentes formatos de fechas a 'DD/MM/AA al DD/MM/AA'.
    Si no cumple el formato esperado, muestra un error y devuelve una cadena vacía.
    """
    if not fecha_str or fecha_str.lower() == 'nan':
        messagebox.showerror(
            "Error en formato de fechas",
            "No se han encontrado fechas. Deben seguir el formato DD/MM/AA al DD/MM/AA"
        )
        return -1
    fecha_str = fecha_str.strip()
    # Buscar patrones tipo 'del XX/XX/XX al XX/XX/XX'
    match = re.search(r'(\d{1,2}/\d{1,2}/\d{2,4})\s*(?:-|al|a|hasta)\s*(\d{1,2}/\d{1,2}/\d{2,4})', fecha_str, re.IGNORECASE)
    if match:
        f1, f2 = match.group(1), match.group(2)
    else:
        # Buscar patrón 'XX/XX/XX-XX/XX/XX'
        match = re.search(r'(\d{1,2}/\d{1,2}/\d{2,4})\s*-\s*(\d{1,2}/\d{1,2}/\d{2,4})', fecha_str)
        if match:
            f1, f2 = match.group(1), match.group(2)
        else:
            # Buscar dos fechas separadas por espacio
            fechas = re.findall(r'\d{1,2}/\d{1,2}/\d{2,4}', fecha_str)
            if len(fechas) >= 2:
                f1, f2 = fechas[0], fechas[1]
            else:
                # Si solo hay una fecha, repetirla
                if fechas:
                    f1 = f2 = fechas[0]
                else:
                    # No se reconoce el formato, mostrar error y devolver cadena vacía
                    messagebox.showerror(
                        "Error en formato de fechas",
                        "Las fechas no están bien, deben seguir el formato DD/MM/AA al DD/MM/AA"
                    )
                    return -1
    # Formatear fechas a DD/MM/AA
    def corta_fecha(f):
        try:
            dt = datetime.strptime(f, "%d/%m/%Y")
            return dt.strftime("%d/%m/%y")
        except Exception:
            try:
                dt = datetime.strptime(f, "%d/%m/%y")
                return dt.strftime("%d/%m/%y")
            except Exception:
                return f
    return f"{corta_fecha(f1)} al {corta_fecha(f2)}"

def extraer_datos_identificativos(nombre_archivo):
    df = pd.read_excel(nombre_archivo, header=None)
    etiquetas = [
        'CÓDIGO EDICIÓN / CODI EDICIÓ',
        'TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA',
        'FECHAS REALIZACIÓN / DATES REALITZACIÓ',
        'MODALIDAD/MODALITAT'
    ]
    resultado = {}
    for i, row in df.iterrows():
        celda_c = str(row[2]).strip() if pd.notna(row[2]) else ''
        for etiqueta in etiquetas:
            if etiqueta.lower() in celda_c.lower():
                valor = str(row[3]).strip() if pd.notna(row[3]) else ''
                # Normalizar fechas si es el campo de fechas
                if etiqueta == 'FECHAS REALIZACIÓN / DATES REALITZACIÓ':
                    valor = normaliza_fechas_realizacion(valor)
                    if valor == -1:
                        return -1
                resultado[etiqueta] = valor
    #print(json.dumps(resultado, ensure_ascii=False, indent=2))
    return resultado

# Ejemplo de uso:
# extraer_datos_identificativos('FITXA-ECONOMICA.xlsx')


def process_excel(nombre_archivo, status_label):
    status_label.config(text="Procesando archivo Excel...")
    xl = pd.ExcelFile(nombre_archivo)
    hoja = xl.sheet_names[0]
    df = xl.parse(hoja, header=None)
    cabecera = [
        'NOMBRE Y APELLIDOS o EMPRESA / NOM I COGNOMS o EMPRESA',
        'DNI / CIF',
        'JURÍDICO',
        'MINUTA / DIETA / FACTURA/ MATERIAL',
        "TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*",
        'UNIDADES/UNITATS',
        'Solo en caso de tutorización indicar Nº ALUMNANOS/AS TUTORIZADOS',
        'TARIFICACIÓN APLICADA (€)',
        'IMPORTE / IMPORT (€)'
    ]
    # Fila 22 (índice 21), y reseteamos índices
    data_rows = df.iloc[21:, :].reset_index(drop=True)
    header_row = data_rows.iloc[0].fillna('')
    # Detección automática de las columnas de interés
    mapping = {}
    for col in range(len(header_row)):
        cell = str(header_row[col]).strip()
        for campo in cabecera:
            if campo in cell:
                mapping[col] = campo
                break
    datos = []
    for i in range(1, len(data_rows)):
        fila = data_rows.iloc[i]
        entry = {}
        for col, campo in mapping.items():
            entry[campo] = fila[col]
        nombre = str(entry.get(cabecera[0], '')).strip()
        dni = str(entry.get(cabecera[1], '')).strip()
        # Solo añadimos filas con nombre y dni válidos
        if nombre and dni and nombre.lower() != 'nan' and dni.lower() != 'nan':
            datos.append(entry)
    # Agrupamos por persona/empresa
    agrupado = {}
    for entry in datos:
        nombre = str(entry.get(cabecera[0], '')).strip()
        dni = str(entry.get(cabecera[1], '')).strip()
        clave = (nombre, dni)
        if clave not in agrupado:
            agrupado[clave] = []
        # Excluimos campos de nombre y dni internos por síntesis
        entry_light = {k: v for k, v in entry.items() if k not in [cabecera[0], cabecera[1]]}
        agrupado[clave].append(entry_light)
    # Montamos el JSON deseado como lista de personas
    resultado = []
    for (nombre, dni), movimientos in agrupado.items():
        resultado.append({
            "Nombre": nombre,
            "DNI": dni,
            "Movimientos": movimientos
        })
    # Mostrar por pantalla
    print(json.dumps(resultado, ensure_ascii=False, indent=2))
    return resultado



def generar_certificas(datos, identificativos, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)



    imagen_path = resource_path('a.png')
    
    # imagen_path = "./a.png"
    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(13)

    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    encabezado.add_run("\nJORDI MARTÍ GUIU, SUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT\n \n")

    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False

    fila = tabla.rows[0]
    fila.cells[0].width = Cm(3)
    fila.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    run = fila.cells[0].paragraphs[0].add_run("CERTIFICA")
    run.bold = True

    movimientos = datos['Movimientos']
    dni = str(datos['DNI'])
    nombre = datos['Nombre'].upper()

    partes = []
    for mov in movimientos:
        if mov.get('TARIFICACIÓN APLICADA (€)', 0) == 300:
            tipo_intervencio = mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "").strip().lower()
            if tipo_intervencio == "tutorización":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "tutorització"
                pagament_text = " setmanes a 300 euros/setmana"
            elif tipo_intervencio == "elaboración de casos-actividades prácticas":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "el·laboració de casos-activitats pràctiques"
                pagament_text = " unitats a 300 euros/unitat"
            elif tipo_intervencio == "ponente":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "ponent"
                pagament_text = " hores a 90 euros/hora"
            else:
                pagament_text = " REVISAR TIPO INTERVENCIÓ"
        else:
            pagament_text = f" hores a {str(mov.get('TARIFICACIÓN APLICADA (€)', '')).lower()} euros/hora"
        partes.append(str(mov.get('UNIDADES/UNITATS', '')).lower() + pagament_text + " en concepte de " + str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).lower())

    cadena_movs = ", ".join(partes)
    # Determinar el código de aplicación según el campo 'JURÍDICO'
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        aplicacion = "233.02"
        juridico = "funcionari"
    else:
        aplicacion = "226.06"

    # Calcular la suma total de los importes
    total_importe = sum(
        float(mov.get('IMPORTE / IMPORT (€)', 0) or 0)
        for mov in movimientos
        if mov.get('IMPORTE / IMPORT (€)', '') not in [None, '', 'nan']
    )

    # Añadir el cuerpo con palabras clave en negrita
    parrafo_cuerpo = fila.cells[1].paragraphs[0]


    def add_bold(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True

    def add_normal(paragraph, text):
        paragraph.add_run(text)

    # Construir el cuerpo con formato
    add_normal(parrafo_cuerpo, f"{nombre} amb NIF: {dni.upper()}\n")
    add_normal(parrafo_cuerpo, juridico + ", Designat/da per esta Subdirecció, ha impartit satisfactòriament " + cadena_movs + " del curs amb les següents dades:\n\n")

    if identificativos:
        add_bold(parrafo_cuerpo, "Codi: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Títol: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Data de realització: ")
        add_normal(parrafo_cuerpo, str("del " + identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')).lower() + "\n")
        add_bold(parrafo_cuerpo, "Lloc de realització: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('MODALIDAD/MODALITAT', '')).upper() + "\n\n")

    add_normal(parrafo_cuerpo, "Per la qual cosa, cal fer-li el pagament corresponent per un total de ")
    add_normal(parrafo_cuerpo, str(total_importe) + " euros (")
    add_normal(parrafo_cuerpo, num2words(total_importe, lang='ca').lower())
    add_normal(parrafo_cuerpo, f"), per l’aplicació {aplicacion}")

    fila.cells[1].width = Cm(13)
    # fila.cells[1].text = 
    doc_name = f"{identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')}_CERTIFICA_{datos['Nombre'].replace(' ', '_')}.docx"


    save_path = doc_name
    if sys.platform == 'darwin':
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        initial_dir = os.getcwd()
        # initial_dir = os.getcwd()
        # Extract filename from doc_name, assuming doc_name might include a path
        initial_file = doc_name.split('/')[-1].split('\\')[-1]
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            initialfile=initial_file,
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:  # Only update save_path if user selected a file
            save_path = file_path
        root.destroy() # Destroy the Tkinter root window

    doc.save(save_path)
    # Convertir a PDF si se desea
    '''
    if convertir_pdf_var.get():
        try:
            convert("./" + doc_name, stdout=open(os.devnull, 'w'), stderr=open(os.devnull, 'w'))
        except Exception as e:
            print(f"Error al convertir a PDF: {e}")
    '''
    try:
        messagebox.showinfo("Documento generado", f"✅ Se ha creado '{identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')}_{datos['Nombre'].replace(' ', '_')}.docx' correctamente.")
    except Exception:
        print("Documento generado correctamente")


# Generar documento Word
# datos es un diccionario con 'Nombre', 'DNI' y 'Movimientos' (lista de dicts)

def generar_documento(datos, identificativos, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1)



    imagen_path = resource_path('a.png')
    
    # imagen_path = "./a.png"
    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(13)

    encabezado = doc.add_paragraph()
    encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    encabezado.add_run("\nJORDI MARTÍ GUIU, SUBDIRECTOR GENERAL DE FORMACIÓ DEL PROFESSORAT\n \n")

    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = False

    fila = tabla.rows[0]
    fila.cells[0].width = Cm(3)
    fila.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    run = fila.cells[0].paragraphs[0].add_run("DESIGNA")
    run.bold = True

    movimientos = datos['Movimientos']
    dni = str(datos['DNI'])
    nombre = datos['Nombre'].upper()

    partes = []
    for mov in movimientos:
        if mov.get('TARIFICACIÓN APLICADA (€)', 0) == 300:
            tipo_intervencio = mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "").strip().lower()
            if tipo_intervencio == "tutorización":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "tutorització"
                pagament_text = " setmanes a 300 euros/setmana"
            elif tipo_intervencio == "elaboración de casos-actividades prácticas":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "el·laboració de casos-activitats pràctiques"
                pagament_text = " unitats a 300 euros/unitat"
            elif tipo_intervencio == "ponente":
                mov["TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*"] = "ponent"
                pagament_text = " hores a 90 euros/hora"
            else:
                pagament_text = " REVISAR TIPO INTERVENCIÓ"
        else:
            pagament_text = f" hores a {str(mov.get('TARIFICACIÓN APLICADA (€)', '')).lower()} euros/hora"
        partes.append(str(mov.get('UNIDADES/UNITATS', '')).lower() + pagament_text + " en concepte de " + str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).lower())

    cadena_movs = ", ".join(partes)
    # Determinar el código de aplicación según el campo 'JURÍDICO'
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        aplicacion = "233.02"
        juridico = "funcionari"
    else:
        aplicacion = "226.06"

    # Calcular la suma total de los importes
    total_importe = sum(
        float(mov.get('IMPORTE / IMPORT (€)', 0) or 0)
        for mov in movimientos
        if mov.get('IMPORTE / IMPORT (€)', '') not in [None, '', 'nan']
    )

    cuerpo = (
        f"{nombre} amb nif: {dni.upper()}\n" +
        juridico + ", perquè impartisca " + cadena_movs + " del curs amb les següents dades:\n"
        + ("Codi: " + str(identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')).upper() + "\n" if identificativos else "")
        + ("Títol: " + str(identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')).upper() + "\n\n" if identificativos else "")
        + ("Data de realització: " + str("del " + identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')).lower() + "\n" if identificativos else "")
        + ("Lloc de realització: " + str(identificativos.get('MODALIDAD/MODALITAT', '')).upper() + "\n\n" if identificativos else "")
        + "caldrà pagar-li en concepte d´assistències la quantitat de "
        + str(total_importe) + " euros ("
        + numero_a_letras(total_importe).lower() +
        f"), per l’aplicació {aplicacion}"
    )


    # Añadir el cuerpo con palabras clave en negrita
    parrafo_cuerpo = fila.cells[1].paragraphs[0]


    def add_bold(paragraph, text):
        run = paragraph.add_run(text)
        run.bold = True

    def add_normal(paragraph, text):
        paragraph.add_run(text)

    # Construir el cuerpo con formato
    add_normal(parrafo_cuerpo, f"{nombre} amb NIF: {dni.upper()}\n")
    add_normal(parrafo_cuerpo, juridico + ", perquè impartisca " + cadena_movs + " del curs amb les següents dades:\n\n")

    if identificativos:
        add_bold(parrafo_cuerpo, "Codi: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Títol: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')).upper() + "\n")
        add_bold(parrafo_cuerpo, "Data de realització: ")
        add_normal(parrafo_cuerpo, str("del " + identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')).lower() + "\n")
        add_bold(parrafo_cuerpo, "Lloc de realització: ")
        add_normal(parrafo_cuerpo, str(identificativos.get('MODALIDAD/MODALITAT', '')).upper() + "\n\n")

    add_normal(parrafo_cuerpo, "caldrà pagar-li en concepte d´assistències la quantitat de ")
    add_normal(parrafo_cuerpo, str(total_importe) + " euros (")
    add_normal(parrafo_cuerpo, num2words(total_importe, lang='ca').lower())
    add_normal(parrafo_cuerpo, f"), per l’aplicació {aplicacion}")

    fila.cells[1].width = Cm(13)
    # fila.cells[1].text = 
    doc_name = f"{identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')}_DESIGNA_{datos['Nombre'].replace(' ', '_')}.docx"

    save_path = doc_name
    if sys.platform == 'darwin':
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        initial_dir = os.getcwd()
        # Extract filename from doc_name, assuming doc_name might include a path
        initial_file = doc_name.split('/')[-1].split('\\')[-1]
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            initialfile=initial_file,
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:  # Only update save_path if user selected a file
            save_path = file_path
        root.destroy() # Destroy the Tkinter root window

    doc.save(save_path)
    # Convertir a PDF si se desea
    '''
    if convertir_pdf_var.get():
        try:
            convert("./" + doc_name, stdout=open(os.devnull, 'w'), stderr=open(os.devnull, 'w'))
        except Exception as e:
            print(f"Error al convertir a PDF: {e}")
    '''
    try:
        messagebox.showinfo("Documento generado", f"✅ Se ha creado '{identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')}_{datos['Nombre'].replace(' ', '_')}.docx' correctamente.")
    except Exception:
        print("Documento generado correctamente")

# Ejemplo de uso:
# for persona in datos_json:  # Si tienes una lista de personas
#     generar_documento(persona)


def generar_skills(datos, identificativos, partida, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    imagen_path = resource_path('b.png')
    
    # imagen_path = "./b.png"

    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(10)

    # Encabezado de autoridad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nMarta Armendia Santos, directora general de Formació Professional, de la Conselleria d’Educació, Cultura, Universitats i Ocupació\n").bold = True

    # RESOLUCIÓN
    p_resolc = doc.add_paragraph()
    p_resolc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_resolc = p_resolc.add_run("RESOLC")
    run_resolc.bold = True

    # 1. Designación del personal docente
    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')

    fechas = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')
    modalidad = identificativos.get('MODALIDAD/MODALITAT', '')




    # Determinar si es funcionario GVA
    movimientos = datos['Movimientos']
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        # Si modalidad contiene "online" o "on line" (ignorando mayúsculas/minúsculas), poner "de forma online", si no, poner "presencial a"
        modalidad_lower = modalidad.lower()
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"1. Designar el personal docent que a continuació es relaciona com a formadors, "
            f"per a formar part de l’equip docent que impartirà la formació {codigo} - {curso}, "
            f"{modalidad_text} del {fechas}."
        )
    else:
        modalidad_lower = modalidad.lower()
        nombre = datos.get('Nombre', '')
        dni = datos.get('DNI', '').replace(' ', '')
        # Calcular total de horas (sumar UNIDADES/UNITATS si es relevante)
        total_hores = sum(
            float(mov.get('UNIDADES/UNITATS', 0) or 0)
            for mov in movimientos
            if str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower() == "síncrona"
        )
        # Si no hay horas, dejarlo vacío o poner el total de unidades
        if not total_hores:
            total_hores = sum(float(mov.get('UNIDADES/UNITATS', 0) or 0) for mov in movimientos)
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"
        designa_text = (
            f"1. Designar a les persones que a continuació es relaciona com a formadors, "
            f"per a formar part de l’equip que impartirà la formació {codigo} - {curso}, "
            f"{modalidad_text} del {fechas}."
        )

    doc.add_paragraph(designa_text)

    # TABLA CENTRAL
    movimientos = datos['Movimientos']


    tabla = doc.add_table(rows=1, cols=6)
    # Centrar contenido horizontal y verticalmente en toda la tabla
    # Centrar contenido horizontal y verticalmente en toda la tabla (todas las filas y columnas)
    
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    headers = ["NOM I COGNOMS", "DNI", "UNITATS", "CONCEPTE", "IMPORT PER UNITAT", "TOTAL"]
    for i in range(6):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fondo gris claro para la cabecera
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)

    for mov in movimientos:
        row = tabla.add_row().cells
        row[0].text = datos['Nombre']
        row[1].text = str(datos['DNI'])
        # Añadir sufijo a 'unitats' según el valor de 'concepte'
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        unitats = str(mov.get('UNIDADES/UNITATS', ''))
        if concepte == "síncrona":
            unitats = f"{unitats} hores"
        elif concepte == "elaboración de casos-actividades prácticas":
            unitats = f"{unitats} casos"
        elif concepte == "tutorización":
            unitats = f"{unitats} setmanes"
        row[2].text = unitats
        row[3].text = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", ""))
        # Añadir sufijo a 'tarificació' según el valor de 'concepte'
        tarificacio = str(mov.get('TARIFICACIÓN APLICADA (€)', ''))
        # Traducción de concepte al valenciano
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
            tarificacio = f"{tarificacio} €/hora"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
            tarificacio = f"{tarificacio} €/cas"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
            tarificacio = f"{tarificacio} €/setmana"
        elif concepte == "ponente":
            concepte_val = "ponent"
            tarificacio = f"{tarificacio} €/hora"
        else:
            concepte_val = concepte
        row[3].text = concepte_val
        row[4].text = tarificacio
        row[5].text = str(mov.get('IMPORTE / IMPORT (€)', '')) + " €"

    # Total general
    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)
    row = tabla.add_row().cells
    # Deja las columnas 0 a 4 vacías y pon "TOTAL" en la columna 5
    for i in range(4):
        row[i].text = ""
        # Quitar borde izquierdo y borde inferior
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        # Eliminar borde izquierdo
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:left w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde inferior
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde derecho
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w')))
    
    
    p_total = row[4].paragraphs[0]
    run_total = p_total.add_run("TOTAL")
    run_total.bold = True
    run_total.font.size = Pt(11)
    # Fondo verde claro para la celda "TOTAL"
    tc_total = row[4]._tc
    tcPr_total = tc_total.get_or_add_tcPr()
    shd_total = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_total.append(shd_total)

    p_importe = row[5].paragraphs[0]
    run_importe = p_importe.add_run(str(importe_total)+"€")
    run_importe.bold = True
    run_importe.font.size = Pt(11)
    # Fondo verde claro para la celda de importe total
    tc_importe = row[5]._tc
    tcPr_importe = tc_importe.get_or_add_tcPr()
    shd_importe = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_importe.append(shd_importe)
    
    # Ajustar la altura de la última fila a 2 cm

    tabla.rows[-1].height = Cm(0.8)



    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

    # Traducir y unir los conceptos en valenciano
    conceptos_valenciano = []
    for mov in movimientos:
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
        elif concepte == "ponente":
            concepte_val = "ponent"
        else:
            concepte_val = concepte
        conceptos_valenciano.append(concepte_val)
        # G01090205GE00000.422C00.TE22000053
    doc.add_paragraph(
        "\n2. Aprovar el gasto per un import total de " +
        f"{importe_total} € en concepte de " +
        " i ".join(set(conceptos_valenciano)) +
        ", per la seua participació en l’activitat esmentada i per actuar fora de l’horari normal de treball. "
        f"Este import s’abonarà d’acord amb el Decret 24/1997, d’11 de febrer, i les seues modificacions posteriors, sobre indemnitzacions per raó del servei i gratificacions per serveis extraordinaris, amb càrrec a l’aplicació pressupostària {partida}, del pressupost de la Generalitat Valenciana per a l’any 2025."
    )

    doc.add_paragraph(
        "3. Esta actuació està cofinançada pel Fons Social Europeu i pel Ministeri d’Educació, "
        "Formació Professional i Esports en el marc del programa d’ocupació, formació i educació del període 2021-2027.\n"
    )

    # Pie y firma
    # Pie y firma centrados
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("València, en data i signatura electrònica\n\n\n\n\n")

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("Marta Armendia Santos\nDirectora General de Formació Professional, de la Conselleria d’Educació, Cultura, Universitats i Ocupació")

    # Nombre de archivo y guardado
    doc_name = f"{codigo}_DESIGNA_{datos['Nombre'].replace(' ', '_')}.docx"
    save_path = doc_name
    if sys.platform == 'darwin':
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        initial_dir = os.getcwd()
        # Extract filename from doc_name, assuming doc_name might include a path
        initial_file = doc_name.split('/')[-1].split('\\')[-1]
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            initialfile=initial_file,
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:  # Only update save_path if user selected a file
            save_path = file_path
        root.destroy() # Destroy the Tkinter root window

    doc.save(save_path)



def generar_skills_resolc(datos, identificativos, partida, fecha, centre_educatiu, carrec, numero_a_letras=lambda x:str(x)):

    

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    imagen_path = resource_path('c.png')
    
    # imagen_path = "./b.png"

    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Times New Roman'
    fuente.size = Pt(12)

    # Encabezado de autoridad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nMARTA ARMENDIA SANTOS, DIRECTORA GENERAL DE FORMACIÓ PROFESSIONAL\n")

    # 1. Designación del personal docente
    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')

    fechas = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')
    modalidad = identificativos.get('MODALIDAD/MODALITAT', '')



    # Determinar si es funcionario GVA
    movimientos = datos['Movimientos']
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        # Si modalidad contiene "online" o "on line" (ignorando mayúsculas/minúsculas), poner "de forma online", si no, poner "presencial a"
        modalidad_lower = modalidad.lower()
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial"

        designa_text = (
            f"Vist l'informe del cap de servei del {fecha}, corresponent a la formació {codigo} - {curso} "
            f"realitzada {modalidad_text} del {fechas}.\n"
            f"Vist que els professors han realitzat en els termes establits i de manera adequada la labor "
            f"per a la qual van ser designats."
        )

    else:
        modalidad_lower = modalidad.lower()
        nombre = datos.get('Nombre', '')
        dni = datos.get('DNI', '').replace(' ', '')
        # Calcular total de horas (sumar UNIDADES/UNITATS si es relevante)
        total_hores = sum(
            float(mov.get('UNIDADES/UNITATS', 0) or 0)
            for mov in movimientos
            if str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower() == "síncrona"
        )
        # Si no hay horas, dejarlo vacío o poner el total de unidades
        if not total_hores:
            total_hores = sum(float(mov.get('UNIDADES/UNITATS', 0) or 0) for mov in movimientos)
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial a {modalidad}"
        designa_text = (
            f"Vist l'informe de la {carrec} del {fecha}, corresponent a la formació {codigo} - {curso} "
            f"realitzada {modalidad_text} del {fechas}.\n"
            f"Vist que els professors han realitzat en els termes establits i de manera adequada la labor "
            f"per a la qual van ser designats."
        )

    p1 = doc.add_paragraph(designa_text)

    # RESOLUCIÓN
    p_resolc = doc.add_paragraph()
    p_resolc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_resolc = p_resolc.add_run("RESOLC")
    run_resolc.bold = True

    # TABLA CENTRAL
    movimientos = datos['Movimientos']
    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)
    texto = (
    f"Que ordene el pagament als professors relacionats a continuació, "   
    f"l'import total de {importe_total} €, amb la distribució indicada, per actuar com a "
    f"col·laboradors en l'activitat de formació {codigo} - {curso}, per actuar fora de l'horari normal de treball i amb càrrec a "
    f"l'aplicació pressupostària {partida}, de conformitat amb el DECRET 80/2025, de 3 de juny, del Consell"
    f"sobre indemnitzacions per raó del servei i gratificacions per serveis extraordinaris."
    )

    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    tabla = doc.add_table(rows=1, cols=8)
    # Centrar contenido horizontal y verticalmente en toda la tabla
    # Centrar contenido horizontal y verticalmente en toda la tabla (todas las filas y columnas)
    
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    # Ajuste de altura y alineación de la fila de cabecera, y tamaño de fuente 10
    headers = ["NOM I COGNOMS", "DNI","CENTRE EDUCATIU","CÀRREC", "UNITATS", "CONCEPTE", "IMPORT PER UNITAT", "TOTAL"]
    
    for i in range(8):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        run.font.size = Pt(8)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fondo gris claro para la cabecera
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)

    for mov in movimientos:
        row = tabla.add_row().cells
        # Ajustar el tamaño de la fuente por defecto a 10 pt para el contenido de las filas
        row[0].text = datos['Nombre']
        row[1].text = str(datos['DNI'])
        row[2].text = centre_educatiu
        row[3].text = carrec
        # Añadir sufijo a 'unitats' según el valor de 'concepte'
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        unitats = str(mov.get('UNIDADES/UNITATS', ''))
        if concepte == "síncrona":
            unitats = f"{unitats} hores"
        elif concepte == "elaboración de casos-actividades prácticas":
            unitats = f"{unitats} casos"
        elif concepte == "tutorización":
            unitats = f"{unitats} setmanes"
        row[4].text = unitats
        # row[5].text = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", ""))
        # Añadir sufijo a 'tarificació' según el valor de 'concepte'
        tarificacio = str(mov.get('TARIFICACIÓN APLICADA (€)', ''))
        # Traducción de concepte al valenciano
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
            tarificacio = f"{tarificacio} €/hora"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
            tarificacio = f"{tarificacio} €/cas"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
            tarificacio = f"{tarificacio} €/setmana"
        elif concepte == "ponente":
            concepte_val = "ponent"
            tarificacio = f"{tarificacio} €/hora"
        else:
            concepte_val = concepte
        row[5].text = concepte_val
        row[6].text = tarificacio
        row[7].text = str(mov.get('IMPORTE / IMPORT (€)', '')) + " €"
        for cell in row:  # Recorre TODAS las celdas de la fila
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Total general
    
    row = tabla.add_row().cells
    # Deja las columnas 0 a 4 vacías y pon "TOTAL" en la columna 5
    for i in range(6):
        row[i].text = ""
        # Quitar borde izquierdo y borde inferior
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        # Eliminar borde izquierdo
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:left w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde inferior
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde derecho
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w')))
    
    
    p_total = row[6].paragraphs[0]
    run_total = p_total.add_run("TOTAL")
    run_total.bold = True
    run_total.font.size = Pt(11)
    # Fondo verde claro para la celda "TOTAL"
    tc_total = row[6]._tc
    tcPr_total = tc_total.get_or_add_tcPr()
    shd_total = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_total.append(shd_total)

    p_importe = row[7].paragraphs[0]
    run_importe = p_importe.add_run(str(importe_total)+"€")
    run_importe.bold = True
    run_importe.font.size = Pt(11)
    # Fondo verde claro para la celda de importe total
    tc_importe = row[7]._tc
    tcPr_importe = tc_importe.get_or_add_tcPr()
    shd_importe = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_importe.append(shd_importe)
    
    # Ajustar la altura de la última fila a 2 cm

    tabla.rows[-1].height = Cm(0.8)



    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

    # Traducir y unir los conceptos en valenciano
    conceptos_valenciano = []
    for mov in movimientos:
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
        elif concepte == "ponente":
            concepte_val = "ponent"
        else:
            concepte_val = concepte
        conceptos_valenciano.append(concepte_val)
        # G01090205GE00000.422C00.TE22000053


    


    # Pie y firma
    # Pie y firma centrados
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("\nValència, en data i signatura electrònica\n\n\n\n\n")

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("Marta Armendia Santos\nDirectora General de Formació Professional")

    # Nombre de archivo y guardado
    doc_name = f"{codigo}_RESOLC_{datos['Nombre'].replace(' ', '_')}.docx"
    save_path = doc_name
    if sys.platform == 'darwin':
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        initial_dir = os.getcwd()
        # Extract filename from doc_name, assuming doc_name might include a path
        initial_file = doc_name.split('/')[-1].split('\\')[-1]
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            initialfile=initial_file,
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:  # Only update save_path if user selected a file
            save_path = file_path
        root.destroy() # Destroy the Tkinter root window

    doc.save(save_path)




def generar_skills_certifica(datos, identificativos, numero_a_letras=lambda x:str(x)):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    imagen_path = resource_path('b.png')
    
    # imagen_path = "./b.png"

    doc.add_picture(imagen_path, width=Cm(15.0))

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Calibri'
    fuente.size = Pt(10)

    # Encabezado de autoridad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\nDavid Montalvà Furió, cap del Servei d’Orientació Professional, de la Direcció General de Formació Professional.\n").bold = True

    # RESOLUCIÓN
    p_resolc = doc.add_paragraph()
    p_resolc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_resolc = p_resolc.add_run("INFORMA")
    run_resolc.bold = True

    # 1. Designación del personal docente
    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '')
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')

    fechas = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '')
    modalidad = identificativos.get('MODALIDAD/MODALITAT', '')




    # Determinar si es funcionario GVA
    movimientos = datos['Movimientos']
    juridico = str(movimientos[0].get('JURÍDICO', '')).strip().lower()
    if juridico == "funcionario gva":
        # Si modalidad contiene "online" o "on line" (ignorando mayúsculas/minúsculas), poner "de forma online", si no, poner "presencial a"
        modalidad_lower = modalidad.lower()
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial a {modalidad}"
        designa_text = (
            f"Que el personal docent que es relaciona a continuació ha format part com a personal col·laborador per "
            f"a formar part de l'equip docent que van ser anomenats per resolució de la Direcció General de "
            f"Formació Professional, per al curs «{codigo} - {curso}» "
            f"realitzat {modalidad_text} del {fechas}.\n"
        )
    else:
        modalidad_lower = modalidad.lower()
        nombre = datos.get('Nombre', '')
        dni = datos.get('DNI', '').replace(' ', '')
        # Calcular total de horas (sumar UNIDADES/UNITATS si es relevante)
        total_hores = sum(
            float(mov.get('UNIDADES/UNITATS', 0) or 0)
            for mov in movimientos
            if str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower() == "síncrona"
        )
        # Si no hay horas, dejarlo vacío o poner el total de unidades
        if not total_hores:
            total_hores = sum(float(mov.get('UNIDADES/UNITATS', 0) or 0) for mov in movimientos)
        if "online" in modalidad_lower or "on line" in modalidad_lower or "semipresencial" in modalidad_lower:
            modalidad_text = f"de forma {modalidad_lower}"
        else:
            modalidad_text = f"presencial a {modalidad}"
        designa_text = (
            f"Que el personal que es relaciona a continuació ha format part com a personal col·laborador per "
            f"a formar part de l'equip docent que van ser anomenats per resolució de la Direcció General de "
            f"Formació Professional, per al curs «{codigo} - {curso}» "
            f"realitzat {modalidad_text} del {fechas}.\n"
        )

    doc.add_paragraph(designa_text)

    # TABLA CENTRAL
    movimientos = datos['Movimientos']


    tabla = doc.add_table(rows=1, cols=6)
    # Centrar contenido horizontal y verticalmente en toda la tabla
    # Centrar contenido horizontal y verticalmente en toda la tabla (todas las filas y columnas)
    
    tabla.style = 'Table Grid'
    hdr_cells = tabla.rows[0].cells
    headers = ["NOM I COGNOMS", "DNI", "UNITATS", "CONCEPTE", "IMPORT PER UNITAT", "TOTAL"]
    for i in range(6):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(headers[i])
        run.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Fondo gris claro para la cabecera
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
        tcPr.append(shd)

    for mov in movimientos:
        row = tabla.add_row().cells
        row[0].text = datos['Nombre']
        row[1].text = str(datos['DNI'])
        # Añadir sufijo a 'unitats' según el valor de 'concepte'
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        unitats = str(mov.get('UNIDADES/UNITATS', ''))
        if concepte == "síncrona":
            unitats = f"{unitats} hores"
        elif concepte == "elaboración de casos-actividades prácticas":
            unitats = f"{unitats} unitats"
        elif concepte == "tutorización":
            unitats = f"{unitats} setmanes"
        row[2].text = unitats
        row[3].text = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", ""))
        # Añadir sufijo a 'tarificació' según el valor de 'concepte'
        tarificacio = str(mov.get('TARIFICACIÓN APLICADA (€)', ''))
        # Traducción de concepte al valenciano
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
            tarificacio = f"{tarificacio} €/hora"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
            tarificacio = f"{tarificacio} €/unitat"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
            tarificacio = f"{tarificacio} €/setmana"
        elif concepte == "ponente":
            concepte_val = "ponent"
            tarificacio = f"{tarificacio} €/hora"   
        else:
            concepte_val = concepte
        row[3].text = concepte_val
        row[4].text = tarificacio
        row[5].text = str(mov.get('IMPORTE / IMPORT (€)', '')) + " €"

    # Total general
    importe_total = sum(float(mov.get('IMPORTE / IMPORT (€)', 0) or 0) for mov in movimientos)
    row = tabla.add_row().cells
    # Deja las columnas 0 a 4 vacías y pon "TOTAL" en la columna 5
    for i in range(4):
        row[i].text = ""
        # Quitar borde izquierdo y borde inferior
        tc = row[i]._tc
        tcPr = tc.get_or_add_tcPr()
        # Eliminar borde izquierdo
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:left w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde inferior
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="nil"/></w:tcBorders>' % nsdecls('w')))
        # Eliminar borde derecho
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w')))
    
    
    p_total = row[4].paragraphs[0]
    run_total = p_total.add_run("TOTAL")
    run_total.bold = True
    run_total.font.size = Pt(11)
    # Fondo verde claro para la celda "TOTAL"
    tc_total = row[4]._tc
    tcPr_total = tc_total.get_or_add_tcPr()
    shd_total = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_total.append(shd_total)

    p_importe = row[5].paragraphs[0]
    run_importe = p_importe.add_run(str(importe_total)+"€")
    run_importe.bold = True
    run_importe.font.size = Pt(11)
    # Fondo verde claro para la celda de importe total
    tc_importe = row[5]._tc
    tcPr_importe = tc_importe.get_or_add_tcPr()
    shd_importe = parse_xml(r'<w:shd {} w:fill="C6EFCE"/>'.format(nsdecls('w')))
    tcPr_importe.append(shd_importe)
    
    # Ajustar la altura de la última fila a 2 cm

    tabla.rows[-1].height = Cm(0.8)



    for row in tabla.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    

    # Traducir y unir los conceptos en valenciano
    conceptos_valenciano = []
    for mov in movimientos:
        concepte = str(mov.get("TIPO DE INTERVENCIÓN*/ TIPUS D'INTERVENCIÓ*", "")).strip().lower()
        if concepte == "síncrona":
            concepte_val = "formació síncrona"
        elif concepte == "elaboración de casos-actividades prácticas":
            concepte_val = "el·laboració de casos-activitats pràctiques"
        elif concepte == "tutorización":
            concepte_val = "tutorització"
        else:
            concepte_val = concepte
        conceptos_valenciano.append(concepte_val)

    doc.add_paragraph(
        "\nAquesta actuació està co-finançada per el Fons Social Europeu i per el Ministeri d'Educació, Formació Professional i Esports en el marc del programa d'ocupació, formació i educació 2021-2027.\n "
    )

    # Pie y firma
    # Pie y firma centrados
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.add_run("València, en data i signatura electrònica\n\n\n\n\n")

    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.add_run("David Montalvà Furió\nCap de servei d'Orientació Professional, de la Direcció General de Formació Professional")

    # Nombre de archivo y guardado
    doc_name = f"{codigo}_CERTIFICA_INFORME_POSTERIOR_{datos['Nombre'].replace(' ', '_')}.docx"
    save_path = doc_name
    if sys.platform == 'darwin':
        root = tk.Tk()
        root.withdraw()  # Hide the main window
        initial_dir = os.getcwd()
        # Extract filename from doc_name, assuming doc_name might include a path
        initial_file = doc_name.split('/')[-1].split('\\')[-1]
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            initialfile=initial_file,
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:  # Only update save_path if user selected a file
            save_path = file_path
        root.destroy() # Destroy the Tkinter root window

    doc.save(save_path)

###### GENERA RESOLC I INFORME SKILLS ######





def show_json(json_data):
    top = tk.Toplevel()
    top.title("Datos en JSON")
    text = tk.Text(top, wrap='word', width=100, height=30)
    text.insert('1.0', json_data)
    text.pack(expand=True, fill='both')


def minuta_skills(datos, identificativos, parent=None):
    # datos puede ser lista de personas o una sola persona
    personas = datos if isinstance(datos, list) else [datos]

    def to_float(v):
        try:
            s = str(v).replace("€", "").replace(",", ".").strip()
            return float(s) if s and s.lower() != "nan" else 0.0
        except Exception:
            return 0.0

    curso = identificativos.get('TÍTULO ACCIÓN FORMATIVA / TÍTOL ACCIÓ FORMATIVA', '') or ''
    codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '') or ''
    nombre_curso_prefill = f"{codigo} - {curso}".strip(" -")
    dates = identificativos.get('FECHAS REALIZACIÓN / DATES REALITZACIÓ', '') or ''

    # Ventana
    top = tk.Toplevel(master=parent)
    top.title("Minuta Skills - Datos por persona")
    top.geometry("900x600")

    # Traer al frente y hacerla modal
    try:
        top.lift()
        top.attributes("-topmost", True)
        top.after(200, lambda: top.attributes("-topmost", False))
        top.grab_set()
        top.focus_force()
    except Exception:
        pass

    def on_close():
        try:
            top.grab_release()
        except Exception:
            pass
        top.destroy()

    top.protocol("WM_DELETE_WINDOW", on_close)

    # Scrollable container
    container = tk.Frame(top)
    container.pack(fill="both", expand=True)

    canvas = tk.Canvas(container, highlightthickness=0)
    vsb = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
    hsb = tk.Scrollbar(container, orient="horizontal", command=canvas.xview)
    scroll_frame = tk.Frame(canvas)

    # Crear ventana dentro del canvas
    window_id = canvas.create_window((0, 0), window=scroll_frame, anchor="nw")

    def on_frame_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def on_canvas_configure(event):
        # Ajustar ancho del frame interior al canvas
        canvas.itemconfig(window_id, width=event.width)

    scroll_frame.bind("<Configure>", on_frame_configure)
    canvas.bind("<Configure>", on_canvas_configure)

    # Soporte rueda ratón
    def _on_mousewheel(event):
        if event.delta:
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        else:
            # Linux
            if event.num == 4:
                canvas.yview_scroll(-3, "units")
            elif event.num == 5:
                canvas.yview_scroll(3, "units")

    canvas.bind_all("<MouseWheel>", _on_mousewheel)
    canvas.bind_all("<Button-4>", _on_mousewheel)
    canvas.bind_all("<Button-5>", _on_mousewheel)

    canvas.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
    canvas.pack(side="left", fill="both", expand=True)
    vsb.pack(side="right", fill="y")
    hsb.pack(side="bottom", fill="x")

    entry_vars = []  # Para recolectar luego

    # Campos por persona
    for idx, persona in enumerate(personas, start=1):
        # Solo procesar si todos los movimientos son 'minuta'
        movs = persona.get('Movimientos', [])
        if not all(
            str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta' or str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'caso-actividad'
            for mov in movs
        ):          
            continue
        movs = persona.get("Movimientos", [])
        total = round(sum(to_float(mov.get('IMPORTE / IMPORT (€)', 0)) for mov in movs), 2)
        neto = round(total * 0.85, 2)  # total - 15%
        bruto = round(total, 2)

        lf = tk.LabelFrame(scroll_frame, text=f"Persona {idx}", padx=10, pady=10)
        lf.pack(fill="x", padx=10, pady=8)

        # Preparar variables
        vars_map = {
            "Nombre y Apellidos": tk.StringVar(value=str(persona.get("Nombre", ""))),
            "NIF": tk.StringVar(value=str(persona.get("DNI", ""))),
            "Domicilio": tk.StringVar(value=""),
            "CP": tk.StringVar(value=""),
            "Población": tk.StringVar(value=""),
            "Provincia": tk.StringVar(value=""),
            "Nombre del curso": tk.StringVar(value=nombre_curso_prefill),
            "Importe bruto": tk.StringVar(value=f"{bruto:.2f}"),
            "Importe neto": tk.StringVar(value=f"{neto:.2f}"),
            "IBAN": tk.StringVar(value=""),
            "BIC": tk.StringVar(value=""),
            "Email": tk.StringVar(value=""),
            "Teléfono": tk.StringVar(value=""),
            "Grup": tk.StringVar(value=""),
            "Nivell": tk.StringVar(value=""),
            "Relacio_juridica": tk.StringVar(value="FI"),
            "Dates_inici_final": tk.StringVar(value=dates),
        }

        entry_vars.append(vars_map)

        # Layout en grid 2 columnas de etiquetas/entradas
        labels = list(vars_map.keys())
        for i, label in enumerate(labels):
            r = i // 2
            c = (i % 2) * 2
            tk.Label(lf, text=label + ":").grid(row=r, column=c, sticky="e", padx=5, pady=4)
            e = tk.Entry(lf, textvariable=vars_map[label], width=40)
            e.grid(row=r, column=c + 1, sticky="w", padx=5, pady=4)
            if label == "Relacio_juridica":
                combo = ttk.Combobox(lf, textvariable=vars_map[label], values=["FI", "FC"], state="readonly", width=37)
                combo.grid(row=r, column=c + 1, sticky="w", padx=5, pady=4)
            if label == "Nivell":
                combo = ttk.Combobox(lf, textvariable=vars_map[label], values=["A26", "A24", "No aplica"], state="readonly", width=37)
                combo.current(1)  # Selecciona "A24" por defecto (índice 1)
                combo.grid(row=r, column=c + 1, sticky="w", padx=5, pady=4)


    '''
    def recopilar_datos():
        salida = []
        for vm in entry_vars:
            salida.append({
                "nombre_apellidos": vm["Nombre y Apellidos"].get(),
                "nif": vm["NIF"].get(),
                "domicilio": vm["Domicilio"].get(),
                "cp": vm["CP"].get(),
                "poblacion": vm["Población"].get(),
                "provincia": vm["Provincia"].get(),
                "curso": vm["Nombre del curso"].get(),
                "importe_bruto": vm["Importe bruto"].get(),
                "importe_neto": vm["Importe neto (bruto - 15%)"].get(),
                "iban": vm["IBAN"].get(),
                "bic": vm["BIC"].get(),
            })
        try:
            show_json(json.dumps(salida, ensure_ascii=False, indent=2))
        except Exception:
            print(json.dumps(salida, ensure_ascii=False, indent=2))
    '''
    def recopilar_y_crear():
        datos_recopilados = []
        for vm in entry_vars:
            datos_recopilados.append({
                "Nombre": vm["Nombre y Apellidos"].get(),
                "NIF": vm["NIF"].get(),
                "Domicili": vm["Domicilio"].get(),
                "CP": vm["CP"].get(),
                "Población": vm["Población"].get(),
                "Provincia": vm["Provincia"].get(),
                "Nombre del curso": vm["Nombre del curso"].get(),
                "Importe bruto": vm["Importe bruto"].get(),
                "Importe neto": vm["Importe neto"].get(),
                "IBAN": vm["IBAN"].get(),
                "BIC": vm["BIC"].get(),
                "Email": vm["Email"].get() if "Email" in vm else "",
                "Teléfono": vm["Teléfono"].get() if "Teléfono" in vm else (vm["Telefono"].get() if "Telefono" in vm else ""),
                "Grup": vm["Grup"].get() if "Grup" in vm else "",
                "Nivell": vm["Nivell"].get() if "Nivell" in vm else "",
                "Relacio_juridica": vm["Relacio_juridica"].get() if "Relacio_juridica" in vm else "",
                "Dates_inici_final": vm["Dates_inici_final"].get() if "Dates_inici_final" in vm else "",
            })
        crea_minuta_skills_docx(datos_recopilados, identificativos)
    
    btn_frame = tk.Frame(top)
    btn_frame.pack(fill="x", padx=10, pady=10)
    tk.Button(btn_frame, text="Crear minutas", command=recopilar_y_crear).pack(side="left", padx=5)
    tk.Button(btn_frame, text="Cerrar", command=on_close).pack(side="right", padx=5)




def crea_minuta_skills_docx(dades, identificativos):
    def crea_docx(datos):
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.1)
        section.footer_distance = Cm(0.1)

        imagen_path = resource_path('a.png')
        
        # imagen_path = "./a.png"
        doc.add_picture(imagen_path, width=Cm(15.0))

        estilo = doc.styles['Normal']
        fuente = estilo.font
        fuente.name = 'Calibri'
        fuente.size = Pt(13)

        encabezado = doc.add_paragraph()
        encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
        encabezado.add_run("GRATIFICACIÓ PER ACTIVITAT DOCENT")
        encabezado.runs[0].bold = True
        encabezado.runs[0].font.size = Pt(14)

            # TABLA DE DATOS DEL PERCEPTOR/A
        tabla = doc.add_table(rows=8, cols=6)
        tabla.autofit = True
        for row in tabla.rows:
            row.height = Cm(0.8)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Borde externo doble y sin bordes internos

        tbl = tabla._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.append(tblPr)

        borders = parse_xml(r'''
            <w:tblBorders %s>
            <w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:right w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)

        # Unir las celdas de la primera fila desde la segunda hasta la última
        primera_fila = tabla.rows[0]
        celda_merged = primera_fila.cells[0]
        for i in range(1, len(primera_fila.cells)):
            celda_merged = celda_merged.merge(primera_fila.cells[i])

        run = primera_fila.cells[0].paragraphs[0].add_run("DADES DEL PERCEPTOR/A")
        run.underline = True
        run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        primera_fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       

        segona_fila = tabla.rows[1]
        celda_merged = segona_fila.cells[0]
        for i in range(1, len(segona_fila.cells)):
            celda_merged = celda_merged.merge(segona_fila.cells[i])

        

        segona_fila.height = Cm(0.2)
        segona_fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        

        tercera_fila = tabla.rows[2]
        celda_merged = tercera_fila.cells[0]
        for i in range(1, len(tercera_fila.cells)):
            celda_merged = celda_merged.merge(tercera_fila.cells[i])
        run2 = tercera_fila.cells[0].paragraphs[0].add_run("NOM I COGNOMS: ")
        run2.bold = True
        run3 = tercera_fila.cells[0].paragraphs[0].add_run(str(datos.get("Nombre", "")))
        run3.bold = False


        cuarta_fila = tabla.rows[3]
        celda_merged = cuarta_fila.cells[0]
        for i in range(1, len(cuarta_fila.cells)):
            celda_merged = celda_merged.merge(cuarta_fila.cells[i])

        cuarta_fila.height = Cm(0.2)
        cuarta_fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


        quinta_fila = tabla.rows[4]
        # NIF | <nif> | Email | <email> | Telèfon | <tel>
        #celda_merged = quinta_fila.cells[0].merge(quinta_fila.cells[1])
        celda_merged = quinta_fila.cells[0]
        for i in range(1, len(quinta_fila.cells)):
            celda_merged = celda_merged.merge(quinta_fila.cells[i])


        run_nif = quinta_fila.cells[0].paragraphs[0].add_run("NIF: ")
        run_nif.bold = True
        run_nif2 = quinta_fila.cells[0].paragraphs[0].add_run(str(datos.get("NIF", datos.get("DNI", datos.get("NIF/NIE", ""))))+"  ")
        run_nif2.bold = False
        
        run_email = quinta_fila.cells[0].paragraphs[0].add_run("Email: ")
        run_email.bold = True
        run_email2 = quinta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Email", datos.get("Correo electrónico", datos.get("Correo", "")))).lower() + "  ")
        run_email2.bold = False
        

        run_tel = quinta_fila.cells[0].paragraphs[0].add_run("Telèfon: ")
        run_tel.bold = True
        run_tel2 = quinta_fila.cells[0].paragraphs[0].add_run(str(
            datos.get("Telèfon", datos.get("Teléfono", datos.get("Telefono", "")))
        ))
        run_tel2.bold = False


        sexta_fila = tabla.rows[5]

        celda_merged = sexta_fila.cells[0].merge(sexta_fila.cells[1])
        run_grup = sexta_fila.cells[0].paragraphs[0].add_run("GRUP: ")
        run_grup.bold = True
        run_grup2 = sexta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Grup", "")))
        run_grup2.bold = False


        celda_merged = sexta_fila.cells[2].merge(sexta_fila.cells[3])
        run_nivell = sexta_fila.cells[2].paragraphs[0].add_run("NIVELL: ")
        run_nivell.bold = True
        run_nivell2 = sexta_fila.cells[2].paragraphs[0].add_run(str(datos.get("Nivell", "")))
        run_nivell2.bold = False

        celda_merged = sexta_fila.cells[4].merge(sexta_fila.cells[5])
        run_relacion = sexta_fila.cells[4].paragraphs[0].add_run("RELACIÓ JURÍDICA: ")
        run_relacion.bold = True
        run_relacion2 = sexta_fila.cells[4].paragraphs[0].add_run(str(datos.get("Relacio_juridica", "")))
        run_relacion2.bold = False

        sexta_fila.height = Cm(1.3)
        sexta_fila.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        septima_fila = tabla.rows[6]

        celda_merged = septima_fila.cells[0]
        for i in range(1, len(septima_fila.cells)):
            celda_merged = celda_merged.merge(septima_fila.cells[i])

        
        run_domicilio = septima_fila.cells[0].paragraphs[0].add_run("DOMICILI: ")
        run_domicilio.bold = True
        run_domicilio2 = septima_fila.cells[0].paragraphs[0].add_run(str(datos.get("Domicili", "")))
        run_domicilio2.bold = False


        # Octava fila: CP, POBLACIÓ, PROVÍNCIA

        octava_fila = tabla.rows[7]

        run_cp = octava_fila.cells[0].paragraphs[0].add_run("CP: ")
        run_cp.bold = True
        run_cp2 = octava_fila.cells[0].paragraphs[0].add_run(str(datos.get("CP", "")))
        run_cp2.bold = False

        celda_merged = octava_fila.cells[1].merge(octava_fila.cells[2])
        run_poblacio = octava_fila.cells[1].paragraphs[0].add_run("POBLACIÓ: ")
        run_poblacio.bold = True
        run_poblacio2 = octava_fila.cells[1].paragraphs[0].add_run(str(datos.get("Población", "")))
        run_poblacio2.bold = False

        celda_merged = octava_fila.cells[4].merge(octava_fila.cells[5])
        run_provincia = octava_fila.cells[4].paragraphs[0].add_run("PROVÍNCIA: ")
        run_provincia.bold = True
        run_provincia2 = octava_fila.cells[4].paragraphs[0].add_run(str(datos.get("Provincia", "")))
        run_provincia2.bold = False

        #####

        doc.add_paragraph("")

        ####

        # TABLA DE DATOS ECONÓMICOS DEL PERCEPTOR/A
        tabla3 = doc.add_table(rows=6, cols=6)
        # tabla.autofit = True
        for row in tabla3.rows:
            row.height = Cm(0.8)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Borde externo doble y sin bordes internos

        tbl = tabla3._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.append(tblPr)

        borders = parse_xml(r'''
            <w:tblBorders %s>
            <w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:right w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)

        # Unir las celdas de la primera fila desde la segunda hasta la última
        primera_fila = tabla3.rows[0]
        celda_merged = primera_fila.cells[0]
        for i in range(1, len(primera_fila.cells)):
            celda_merged = celda_merged.merge(primera_fila.cells[i])

        run = primera_fila.cells[0].paragraphs[0].add_run("DADES ECONÒMIQUES")
        run.underline = True
        run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        primera_fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       

        segona_fila = tabla3.rows[1]
        run2 = segona_fila.cells[0].paragraphs[0].add_run("DESCRIPCIÓ DE L’ACTIVITAT: ")
        run2.bold = True

        celda_merged = segona_fila.cells[0]
        for i in range(0, len(segona_fila.cells)):
            celda_merged = celda_merged.merge(segona_fila.cells[i])

        tercera_fila = tabla3.rows[2]
        tercera_fila.cells[1].paragraphs[0].text = datos["Nombre del curso"]

        celda_merged = tercera_fila.cells[0]
        for i in range(0, len(tercera_fila.cells)):
            celda_merged = celda_merged.merge(tercera_fila.cells[i])

        cuarta_fila = tabla3.rows[3]

        for i in range(0, len(cuarta_fila.cells)):
            celda_merged = cuarta_fila.cells[0].merge(cuarta_fila.cells[i])

        cuarta_fila.cells[0].paragraphs[0].add_run("DATES: ")
        cuarta_fila.cells[0].paragraphs[0].runs[0].bold = True
        run2 = cuarta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Dates_inici_final", "")))
        run2.bold = False


        quinta_fila = tabla3.rows[4]
        for i in range(0, len(quinta_fila.cells)):
            celda_merged = quinta_fila.cells[0].merge(quinta_fila.cells[i])

        
        run = quinta_fila.cells[0].paragraphs[0].add_run("IMPORT BRUT: ")
        run.bold = True
        run2 = quinta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Importe bruto", ""))+ " €")
        run2.bold = False

        sexta_fila = tabla3.rows[5]
        for i in range(0, len(sexta_fila.cells)):
            celda_merged = sexta_fila.cells[0].merge(sexta_fila.cells[i])

        
        sexta_fila.cells[0].paragraphs[0].add_run("IMPORT NET: ")
        sexta_fila.cells[0].paragraphs[0].runs[0].bold = True
        run2 = sexta_fila.cells[0].paragraphs[0].add_run(str(datos.get("Importe neto", "")) + " €")
        run2.bold = False
        



        ###########################################

        doc.add_paragraph("")
        

        ####

        # TABLA DE DATOS bancarios DEL PERCEPTOR/A
        tabla4 = doc.add_table(rows=3, cols=6)
        # tabla.autofit = True
        for row in tabla4.rows:
            row.height = Cm(0.9)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        # Borde externo doble y sin bordes internos

        tbl = tabla4._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
            tbl.append(tblPr)

        borders = parse_xml(r'''
            <w:tblBorders %s>
            <w:top w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:left w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:bottom w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:right w:val="double" w:sz="12" w:space="0" w:color="000000"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
            </w:tblBorders>
        ''' % nsdecls('w'))
        tblPr.append(borders)

        # Unir las celdas de la primera fila desde la segunda hasta la última
        primera_fila = tabla4.rows[0]
        celda_merged = primera_fila.cells[0]
        for i in range(1, len(primera_fila.cells)):
            celda_merged = celda_merged.merge(primera_fila.cells[i])

        run = primera_fila.cells[0].paragraphs[0].add_run("DADES BANCÀRIES")
        run.underline = True
        run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        primera_fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
       

        segona_filatt4 = tabla4.rows[1]
        celda_merged2 = segona_filatt4.cells[0]
        for i in range(1, len(segona_filatt4.cells)):
            celda_merged2 = celda_merged2.merge(segona_filatt4.cells[i])
   
        run = segona_filatt4.cells[0].paragraphs[0].add_run("IBAN: ")
        run.bold = True
        
        # Formatear IBAN: separar en grupos de 4 caracteres
        iban_raw = str(datos.get("IBAN", "")).replace(" ", "").upper()
        iban_formatted = " ".join([iban_raw[i:i+4] for i in range(0, len(iban_raw), 4)])
        run2 = segona_filatt4.cells[0].paragraphs[0].add_run(iban_formatted)
        run2.bold = False


        # segona_fila.cells[1].paragraphs[0].text = str(datos.get("IBAN", ""))

        tercera_filatt4 = tabla4.rows[2]
        celda_merged = tercera_filatt4.cells[0]
        for i in range(0, len(tercera_filatt4.cells)):
            celda_merged = tercera_filatt4.cells[0].merge(tercera_filatt4.cells[i])

        run = tercera_filatt4.cells[0].paragraphs[0].add_run("BIC: ")
        run.bold = True
        run2 = tercera_filatt4.cells[0].paragraphs[0].add_run(str(datos.get("BIC", "")))
        run2.bold = False

        ####################################

        doc.add_paragraph("Declare que he realitzat la citada activitat en la data que s'assenyala.")

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("(firma digital)")

        # Pie de página
        footer = doc.sections[0].footer
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_paragraph.add_run("CONSELLERIA D'EDUCACIÓ, CULTURA, UNIVERSITATS I OCUPACIÓ\n")
        run.font.size = Pt(7)
        run = footer_paragraph.add_run("Av. Campanar, 32. 46015 - València. CIF S4611001A\n")
        run.font.size = Pt(7)
        run = footer_paragraph.add_run("DIRECCIÓ GENERAL DE FORMACIÓ PROFESSIONAL")
        run.font.size = Pt(7)



        #####################################
        # fila.cells[1].text = 
        codigo = identificativos.get('CÓDIGO EDICIÓN / CODI EDICIÓ', '')
        doc_name = f"{codigo}_MINUTA_{datos['Nombre'].replace(' ', '_')}.docx"

        #doc.save(doc_name)

        save_path = doc_name
        if sys.platform == 'darwin':
            root = tk.Tk()
            root.withdraw()  # Hide the main window
            initial_dir = os.getcwd()
            # Extract filename from doc_name, assuming doc_name might include a path
            initial_file = doc_name.split('/')[-1].split('\\')[-1]
            file_path = filedialog.asksaveasfilename(
                initialdir=initial_dir,
                initialfile=initial_file,
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")]
            )
            if file_path:  # Only update save_path if user selected a file
                save_path = file_path
            root.destroy() # Destroy the Tkinter root window

        doc.save(save_path)

        try:
            messagebox.showinfo("Documento generado", f"✅ Se ha creado '{doc_name}' correctamente.")
        except Exception:
            print("Documento generado correctamente")
    


    import re

    def validar_datos(campos):
        errores = {}

        # Solo texto (letras y espacios)
        solo_texto = lambda s: bool(re.fullmatch(r"[A-Za-zÁÉÍÓÚáéíóúÑñÜü\s]+", s.strip()))
        # NIF: 8 números y 1 letra al final
        nif_valido = lambda s: bool(re.fullmatch(r"\d{8}[A-Za-z]", s.strip()))
        # CP: 5 cifras
        cp_valido = lambda s: bool(re.fullmatch(r"\d{5}", s.strip()))
        # Solo cifras (números con posible punto decimal)
        solo_cifras = lambda s: bool(re.fullmatch(r"\d+(\.\d+)?", s.strip()))
        # IBAN básico: letra-letra y seguido 2-30 dígitos/letras (validez formal)
        iban_valido = lambda s: bool(re.fullmatch(r"[A-Z]{2}\d{2}[A-Z0-9]{10,30}", s.strip().replace(" ", "").upper()))
        # Email básico
        email_valido = lambda s: bool(re.fullmatch(r"[^@]+@[^@]+\.[^@]+", s.strip()))
        # Teléfono: 9 cifras
        telefono_valido = lambda s: bool(re.fullmatch(r"\d{9}", s.strip()))
        # Texto y números (alfanumérico con espacios)
        texto_numeros = lambda s: bool(re.fullmatch(r"[A-Za-z0-9ÁÉÍÓÚáéíóúÑñÜü\s\-\.\,\;\:\(\)\[\]\{\}\¿\?\¡\!\@\#\$\%\&\*\_\+\=\/\\\|\~]+", s.strip()))





        # Validaciones
        errores["Nombre y Apellidos"] = "" if solo_texto(campos.get("Nombre", "")) else "Solo texto permitido"
        errores["NIF"] = "" if nif_valido(campos.get("NIF", "")) else "Debe tener 8 números y 1 letra"
        errores["Domicili"] = "" if texto_numeros(campos.get("Domicili", "")) else "Solo texto permitido"
        errores["CP"] = "" if cp_valido(campos.get("CP", "")) else "Debe tener 5 cifras"
        errores["Población"] = "" if solo_texto(campos.get("Población", "")) else "Solo texto permitido"
        errores["Provincia"] = "" if solo_texto(campos.get("Provincia", "")) else "Solo texto permitido"
        errores["Nombre del curso"] = "" if texto_numeros(campos.get("Nombre del curso", "")) else "Solo texto permitido"
        errores["Importe bruto"] = "" if solo_cifras(campos.get("Importe bruto", "")) else "Solo cifras válidas"
        errores["Importe neto"] = "" if solo_cifras(campos.get("Importe neto", "")) else "Solo cifras válidas"
        errores["IBAN"] = "" if iban_valido(campos.get("IBAN", "")) else "IBAN inválido"
        # BIC no definido: sin validación (puedes añadir)
        errores["BIC"] = ""
        errores["Email"] = "" if email_valido(campos.get("Email", "")) else "Email inválido"
        errores["Teléfono"] = "" if telefono_valido(campos.get("Teléfono", "")) else "Debe tener 9 cifras"
        # Campos sin validación o con opciones libres
        errores["Grup"] = ""
        errores["Nivell"] = ""
        errores["Relacio_juridica"] = ""
        errores["Dates_inici_final"] = ""

        # Si todos los errores están vacíos, devolver {"ok": True}
        if all(v == "" for v in errores.values()):
            return {"ok": True}
        else:
            return {"ok": False, "errores": errores}

    #for dato in dades:
    #    crea_docx(dato)
    #    messagebox.showinfo("Documento generado", f"✅ Se ha creado la minuta correctamente para {dato.get('Nombre', '')}.")

    ##### temporatlmente deshabilitado validación #####


    
    for dato in dades:
        validar_datos_data = validar_datos(dato)
        errores_text = ""
        if not validar_datos_data.get("ok", False):
            for key, val in validar_datos_data['errores'].items():
                if val:
                    if val != "":
                        errores_text += f"Error en {key}: {val}\n"
            messagebox.showwarning(f"Validación de datos en {dato.get('Nombre', '')} \n", f"Errores de validación:\n\n{errores_text}\nPor favor, corrígelos antes de generar el documento.")
            break
        else:
            crea_docx(dato)
            #messagebox.showinfo("Documento generado", f"✅ Se ha creado la minuta correctamente para {dato.get('Nombre', '')}.")
    



def main():
    # global convertir_pdf_var
    global root
    root = tk.Tk()
    root.title("GENERA DESIGNAS")
    # Centrar la ventana en la pantalla
    window_width = 400
    window_height = 350
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = int((screen_width / 2) - (window_width / 2))
    y = int((screen_height / 2) - (window_height / 2))
    root.geometry(f"{window_width}x{window_height}+{x}+{y}")

    status_label = tk.Label(root, text="Haz clic para procesar", fg="blue", font=("Arial", 12))
    status_label.pack(pady=10)

    # convertir_pdf_var = tk.BooleanVar()
    # chk_convertir_pdf = tk.Checkbutton(root, text="Convertir todos los DOCX a PDF al finalizar", variable=convertir_pdf_var)
    # chk_convertir_pdf.pack(pady=5)
    global es_skills
    es_skills = tk.BooleanVar()
    global es_erasmus
    es_erasmus = tk.BooleanVar()

    chk_es_skills = tk.Checkbutton(root, text="Es Skills", variable=es_skills)
    chk_es_skills.pack(pady=5)
    chk_es_erasmus = tk.Checkbutton(root, text="Es fons ERASMUS", variable=es_erasmus)
    chk_es_erasmus.pack(pady=5)
    '''
    ON PROCESS PRINCIPAL
    '''
    def on_process(tipo, parent=root):
        if tipo == "resolc":
             # Ventana para seleccionar fecha, centre educatiu y carrec
             if not es_skills.get() and not es_erasmus.get():
                messagebox.showerror("Error", "Selecciona una opción 'Es Skills' o 'Es fons ERASMUS' per a generar el RESOLC.")
                return

        fecha = ""
        centre_educatiu = ""
        carrec = ""
        fecha_window = None
        def crea_ventana_fechas(nombre):
            nonlocal fecha_window, fecha, centre_educatiu, carrec
            fecha_window = tk.Toplevel(parent)
            fecha_window.title("Seleccionar fecha")

            # Centrar la ventana en la pantalla
            window_width = 300
            window_height = 300
            screen_width = fecha_window.winfo_screenwidth()
            screen_height = fecha_window.winfo_screenheight()
            x = int((screen_width / 2) - (window_width / 2))
            y = int((screen_height / 2) - (window_height / 2))
            fecha_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

            tk.Label(fecha_window, text=nombre, font=("Arial", 11, "bold")).pack(pady=5)
            tk.Label(fecha_window, text="Selecciona la fecha para la resolución:").pack(pady=10)

            fecha_entry = DateEntry(fecha_window, date_pattern='dd/mm/yyyy')
            fecha_entry.pack(pady=5)

            tk.Label(fecha_window, text="Centre educatiu:").pack(pady=5)
            centre_educatiu_entry = tk.Entry(fecha_window, width=30)
            centre_educatiu_entry.pack(pady=5)

            tk.Label(fecha_window, text="Càrrec:").pack(pady=5)
            carrec_entry = tk.Entry(fecha_window, width=30)
            carrec_entry.pack(pady=5)

            btn_frame = tk.Frame(fecha_window)
            btn_frame.pack(pady=10)
            
            confirmar_btn = tk.Button(btn_frame, text="Confirmar", command=lambda: confirmar_fecha())
            confirmar_btn.pack(side="left", padx=5)
            tk.Button(btn_frame, text="Cancelar", command=fecha_window.destroy).pack(side="left", padx=5)
            


            def confirmar_fecha():
                nonlocal fecha, centre_educatiu, carrec
                fecha = fecha_entry.get_date().strftime('%d/%m/%Y')
                centre_educatiu = centre_educatiu_entry.get()
                carrec = carrec_entry.get()
                fecha_window.destroy()
                


        # crea_ventana_fechas()
        # fecha_window.wait_window()

        # messagebox.showinfo("Fecha seleccionada", f"Fecha seleccionada: {fecha}\nCentre educatiu: {centre_educatiu}\nCàrrec: {carrec}")
        
        global t
        t = tipo
        try:
            excel_file = find_excel_file(status_label)
            if not excel_file:
                messagebox.showerror("Error", "No se encontró ningún archivo Excel en la carpeta.")
                return
            json_data = process_excel(excel_file, status_label)

            hoja_excel = extraer_datos_identificativos(excel_file)

            if hoja_excel == -1:
                status_label.config(text="Error en los datos identificativos. Revisa las fechas.")
                return

            #show_json(hoja_excel)

            if t == "min":
                minuta_skills(datos=json_data, identificativos=hoja_excel, parent=root)
                status_label.config(text="¡Proceso completado!")
                return
            elif es_skills.get():
                if es_erasmus.get():
                    messagebox.showerror("Error", "La opción 'Es fons ERASMUS' no es compatible con 'Es Skills'.")
                    return
                for persona in json_data:
                    # Filtrar movimientos que contienen "minuta" en el campo 'MINUTA / DIETA / FACTURA/ MATERIAL'
                    # Solo generar documento si TODOS los movimientos son 'minuta'
                    # Si algún movimiento es "caso-actividad", lo convertimos a "minuta"
                    for mov in persona.get('Movimientos', []):
                        tipo = str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower()
                        if tipo == 'caso-actividad':
                            mov['MINUTA / DIETA / FACTURA/ MATERIAL'] = 'minuta'
                    if all(
                        str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta'
                        for mov in persona.get('Movimientos', [])
                    ):
                        if t == "des":
                            generar_skills(datos=persona, identificativos=hoja_excel, partida="G01090205GE00000.422C00.22699 fons TE22000053")
                        elif t == "cer":
                            generar_skills_certifica(datos=persona, identificativos=hoja_excel)
                        elif t == "resolc":
                            crea_ventana_fechas(persona.get('Nombre', ''))
                            fecha_window.wait_window()
                            generar_skills_resolc(datos=persona, identificativos=hoja_excel, fecha=fecha, centre_educatiu=centre_educatiu, carrec=carrec, partida="G01090205GE00000.422C00.22699 fons TE22000053")
                    status_label.config(text="¡Proceso completado!")
                status_label.config(text="¡Proceso completado!")
                return
                  # Salir después de generar skills si está seleccionado
            elif es_erasmus.get():
                if es_skills.get():
                    messagebox.showerror("Error", "La opción 'Es fons ERASMUS' no es compatible con 'Es Skills'.")
                    return
                for persona in json_data:
                    # Filtrar movimientos que contienen "minuta" en el campo 'MINUTA / DIETA / FACTURA/ MATERIAL'
                    # Solo generar documento si TODOS los movimientos son 'minuta'
                    # Si algún movimiento es "caso-actividad", lo convertimos a "minuta"
                    for mov in persona.get('Movimientos', []):
                        tipo = str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower()
                        if tipo == 'caso-actividad':
                            mov['MINUTA / DIETA / FACTURA/ MATERIAL'] = 'minuta'
                    if all(
                        str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta'
                        for mov in persona.get('Movimientos', [])
                    ):
                        # generar_documento(datos=persona, identificativos=hoja_excel)
                        if t == "des":
                            generar_skills(datos=persona, identificativos=hoja_excel, partida="G01090205GE00000.422C00.22699 fons OT23000000")
                        elif t == "cer":
                            generar_skills_certifica(datos=persona, identificativos=hoja_excel)
                        elif t == "resolc":
                            crea_ventana_fechas(persona.get('Nombre', ''))
                            fecha_window.wait_window()
                            generar_skills_resolc(datos=persona, identificativos=hoja_excel, fecha=fecha,  centre_educatiu=centre_educatiu, carrec=carrec,partida="G01090205GE00000.422C00.22699 fons OT23000000")
                status_label.config(text="¡Proceso completado!")
                return
                  # Salir después de generar designas si no está seleccionado skills
            elif not es_skills.get():
                for persona in json_data:
                    # Filtrar movimientos que contienen "minuta" en el campo 'MINUTA / DIETA / FACTURA/ MATERIAL'
                    # Solo generar documento si TODOS los movimientos son 'minuta'
                    # Si algún movimiento es "caso-actividad", lo convertimos a "minuta"
                    for mov in persona.get('Movimientos', []):
                        tipo = str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower()
                        if tipo == 'caso-actividad':
                            mov['MINUTA / DIETA / FACTURA/ MATERIAL'] = 'minuta'
                    if all(
                        str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta'
                        for mov in persona.get('Movimientos', [])
                    ):
                        # generar_documento(datos=persona, identificativos=hoja_excel)
                        if t == "des":
                            generar_documento(datos=persona, identificativos=hoja_excel)
                        elif t == "cer":
                            generar_certificas(datos=persona, identificativos=hoja_excel)
                status_label.config(text="¡Proceso completado!")
                return
                  # Salir después de generar designas si no está seleccionado skills
            '''
            for persona in json_data:
                # Filtrar movimientos que contienen "minuta" en el campo 'MINUTA / DIETA / FACTURA/ MATERIAL'
                # Solo generar documento si TODOS los movimientos son 'minuta'
                if all(
                    str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta'
                    for mov in persona.get('Movimientos', [])
                ):
                    generar_documento(datos=persona, identificativos=hoja_excel)
            


            for persona in json_data:
                # Filtrar movimientos que contienen "minuta" en el campo 'MINUTA / DIETA / FACTURA/ MATERIAL'
                # Solo generar documento si TODOS los movimientos son 'minuta'
                if all(
                    str(mov.get('MINUTA / DIETA / FACTURA/ MATERIAL', '')).strip().lower() == 'minuta'
                    for mov in persona.get('Movimientos', [])
                ):
                    generar_skills(datos=persona, identificativos=hoja_excel)

            '''

            # show_json(json_data)
            messagebox.showerror("Error", "Ninguna opción seleccionada.")
        except Exception as e:
            status_label.config(text=f"Error: {e}")
            messagebox.showerror("Error", f"Error procesando el archivo: {e}")

    btn = tk.Button(root, text="Genera Designas", command=lambda: on_process("des"), font=("Arial", 12))
    btn2 = tk.Button(root, text="Genera Certifica", command=lambda: on_process("cer"), font=("Arial", 12))
    btn3 = tk.Button(root, text="Genera Minuta DGFP", command=lambda: on_process("min", root), font=("Arial", 12))
    btn4 = tk.Button(root, text="Genera Resolc DGFP", command=lambda: on_process("resolc", root), font=("Arial", 12))
    btn.pack(pady=10)
    btn2.pack(pady=10)
    btn3.pack(pady=10)
    btn4.pack(pady=10)

    version_label = tk.Label(root, text=version, font=("Arial", 10), fg="gray")
    version_label.place(relx=1.0, rely=1.0, anchor='se', x=-5, y=-5)


    root.mainloop()

if __name__ == "__main__":
    main()